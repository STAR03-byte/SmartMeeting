from datetime import date

from docx import Document


def add_section_title(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def add_sub_title(doc: Document, title: str) -> None:
    doc.add_heading(title, level=2)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for idx, text in enumerate(rows[0]):
        hdr[idx].text = text
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text


def main() -> None:
    doc = Document()

    doc.add_heading("项目立项规划书", 0)
    doc.add_paragraph("项目名称：SmartMeeting 智能会议系统")
    doc.add_paragraph("项目类型：Web应用 / SaaS")
    doc.add_paragraph("文档版本：V1.0")
    doc.add_paragraph(f"编制日期：{date.today().isoformat()}")
    doc.add_page_break()

    doc.add_heading("目录", level=1)
    toc = [
        "一、项目基本信息",
        "二、项目概述",
        "三、需求分析",
        "四、产品功能规划",
        "五、目标用户分析",
        "六、运营模式与商业计划",
        "七、技术架构设计",
        "八、团队组织架构",
        "九、项目实施计划",
        "十、项目里程碑",
        "十一、风险评估与应对措施",
        "十二、项目预算",
        "十三、项目交付清单",
        "十四、项目评审标准",
        "十五、附录",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    add_section_title(doc, "一、项目基本信息")
    add_table(
        doc,
        [
            ["项目名称", "SmartMeeting 智能会议系统"],
            ["项目英文名", "SmartMeeting"],
            ["项目类型", "AI + 会议协作 Web 应用"],
            ["开发周期", "6周"],
            ["项目负责人", "star03（独立开发）"],
            ["开发团队", "1人（产品+前后端+测试+运维）"],
            ["文档版本", "V1.0"],
            ["编制日期", date.today().isoformat()],
        ],
    )

    add_section_title(doc, "二、项目概述")
    add_sub_title(doc, "2.1 项目背景")
    doc.add_paragraph(
        "远程协作场景下，会议记录分散、会后任务难追踪。SmartMeeting 以“转写-纪要-任务”闭环为核心，降低会议管理成本。"
    )
    add_sub_title(doc, "2.2 项目简介")
    doc.add_paragraph(
        "SmartMeeting 提供会议管理、音频上传、语音转写、AI纪要与任务抽取能力，并支持任务流转管理，目标在6周内形成可演示MVP。"
    )
    add_sub_title(doc, "2.3 项目目标")
    doc.add_paragraph("短期（1-2周）：搭建后端基础能力与核心数据模型。", style="List Bullet")
    doc.add_paragraph("中期（3-4周）：完成转写与纪要任务闭环，形成端到端可测流程。", style="List Bullet")
    doc.add_paragraph("收官（5-6周）：完成前端联调、文档交付与验收。", style="List Bullet")

    add_section_title(doc, "三、需求分析")
    add_sub_title(doc, "3.1 业务需求")
    for item in [
        "会议信息可结构化沉淀，支持后续复盘。",
        "会后待办自动提取并可追踪状态。",
        "支持单人团队快速迭代，部署与维护成本可控。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    add_sub_title(doc, "3.2 非功能需求")
    add_table(
        doc,
        [
            ["维度", "目标"],
            ["可用性", "核心接口可用性 >= 99%"],
            ["性能", "30分钟会议在3-5分钟内完成纪要输出"],
            ["安全", "敏感配置环境变量化，接口输入校验"],
            ["可维护性", "模块化服务拆分，测试覆盖核心流程"],
        ],
    )

    add_section_title(doc, "四、产品功能规划")
    add_table(
        doc,
        [
            ["模块", "说明", "优先级", "备注"],
            ["会议管理", "会议创建/查询/更新", "P0", "已完成"],
            ["音频上传", "会议音频上传与落盘", "P0", "已完成"],
            ["语音识别", "mock ASR写入转写", "P0", "已完成"],
            ["AI纪要", "摘要生成与版本记录", "P0", "已完成"],
            ["任务提取", "规则抽取与负责人推断", "P0", "已完成"],
            ["任务管理", "状态机流转与筛选", "P0", "已完成"],
            ["前端页面", "总览页/详情页联调", "P1", "已完成"],
        ],
    )

    add_section_title(doc, "五、目标用户分析")
    add_table(
        doc,
        [
            ["角色", "需求", "痛点"],
            ["团队管理者", "快速获知会议结论", "纪要滞后、行动项遗漏"],
            ["项目经理", "任务自动拆分和跟踪", "会后推进依赖人工"],
            ["参会成员", "低成本回顾会议内容", "信息过载"],
        ],
    )

    add_section_title(doc, "六、运营模式与商业计划")
    doc.add_paragraph("MVP阶段以产品可用与种子用户验证为核心，暂不追求商业化收入。", style="List Bullet")
    doc.add_paragraph("后续可采用团队订阅制：按团队规模与AI调用量计费。", style="List Bullet")
    doc.add_paragraph("通过会议场景高频使用形成留存，再扩展高级协作能力。", style="List Bullet")

    add_section_title(doc, "七、技术架构设计")
    add_table(
        doc,
        [
            ["层级", "技术选型"],
            ["前端", "Vue 3 + TypeScript + Element Plus"],
            ["后端", "FastAPI + SQLAlchemy"],
            ["数据库", "MySQL（测试使用SQLite）"],
            ["AI能力", "Whisper/LLM（当前为mock + 规则）"],
            ["部署", "Docker + Nginx（规划中）"],
        ],
    )

    add_section_title(doc, "八、团队组织架构")
    doc.add_paragraph("组织形式：单人全栈交付（产品、开发、测试、文档、部署）。")
    doc.add_paragraph("协作机制：按周复盘，按模块交付，关键功能均以测试验证。")

    add_section_title(doc, "九、项目实施计划")
    add_table(
        doc,
        [
            ["周次", "工作重点", "输出"],
            ["第1周", "后端基础框架与数据模型", "核心API初版"],
            ["第2周", "会议管理与任务管理", "CRUD与测试"],
            ["第3周", "音频上传与转写链路", "上传/转写接口"],
            ["第4周", "纪要与任务抽取", "后处理闭环"],
            ["第5周", "前端页面与联调", "总览/详情可用"],
            ["第6周", "文档、优化、验收", "交付包与演示稿"],
        ],
    )

    add_section_title(doc, "十、项目里程碑")
    doc.add_paragraph("M1（第2周末）：后端可完成会议与任务基础流程。", style="List Bullet")
    doc.add_paragraph("M2（第4周末）：转写、纪要、任务抽取闭环跑通。", style="List Bullet")
    doc.add_paragraph("M3（第6周末）：前后端联调完成，文档齐备可验收。", style="List Bullet")

    add_section_title(doc, "十一、风险评估与应对措施")
    add_table(
        doc,
        [
            ["风险", "影响", "应对"],
            ["单人开发负荷", "进度波动", "严格范围管理，优先P0"],
            ["模型效果不稳定", "纪要质量波动", "先规则兜底，后接LLM"],
            ["联调复杂度", "返工增多", "接口文档先行，测试先行"],
        ],
    )

    add_section_title(doc, "十二、项目预算")
    add_table(
        doc,
        [
            ["成本项", "预估"],
            ["云主机与存储", "300-800 元/月"],
            ["AI调用", "800-3000 元/月（后续）"],
            ["域名与证书", "100-300 元/年"],
            ["总计", "1200-4100 元/月（MVP阶段）"],
        ],
    )

    add_section_title(doc, "十三、项目交付清单")
    for item in [
        "后端代码（会议、转写、任务、后处理）",
        "前端代码（总览页、详情页、联调链路）",
        "测试用例与测试报告",
        "数据库设计文档与接口文档",
        "项目运行手册与README",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_section_title(doc, "十四、项目评审标准")
    for item in [
        "核心P0功能可用并通过测试。",
        "后处理链路可稳定输出摘要与任务。",
        "前后端可完成完整业务演示。",
        "文档完整、可复现部署与运行。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_section_title(doc, "十五、附录")
    doc.add_paragraph("术语：MVP、ASR、LLM、P0/P1/P2。")
    doc.add_paragraph("参考文档：database-design.md、backend-api.md、frontend-runbook.md。")

    output_path = r"D:\乱七八糟\SmartMeeting_项目立项规划书_V1.0.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
