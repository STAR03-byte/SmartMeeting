from datetime import date

from docx import Document


SOURCE_PATH = r"D:\乱七八糟\项目立项规划书.docx"
BACKUP_PATH = r"D:\乱七八糟\项目立项规划书_原始备份.docx"


REPLACEMENTS = {
    "校园学习资源共享平台（CampusShare）": "SmartMeeting 智能会议系统",
    "校园学习资源共享平台": "SmartMeeting 智能会议系统",
    "CampusShare": "SmartMeeting",
    "Campus Learning Resource Sharing Platform": "SmartMeeting Intelligent Meeting System",
    "Web应用/网站": "Web应用/SaaS",
    "8周（2026年3月-2026年5月）": "6周（2026年3月-2026年4月）",
    "软件开发团队": "独立开发项目",
    "[待填写团队名称]": "独立开发（个人项目）",
    "[待填写姓名]": "star03",
    "[待填写学号]": "不适用",
    "[待填写电话]": "不适用",
    "[待填写邮箱]": "star03@example.com",
    "[待填写]": "star03",
    "指导老师": "项目顾问",
    "学习资源": "会议内容",
    "资料": "会议记录",
    "积分": "任务优先级",
}


def replace_text(text: str) -> str:
    updated = text
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    updated = updated.replace("2026年3月10日", f"{date.today().year}年{date.today().month}月{date.today().day}日")
    return updated


def update_paragraphs(paragraphs) -> None:
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.text = replace_text(run.text)


def main() -> None:
    # 先保留一个原始备份
    original = Document(SOURCE_PATH)
    original.save(BACKUP_PATH)

    doc = Document(SOURCE_PATH)
    update_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                update_paragraphs(cell.paragraphs)

    doc.save(SOURCE_PATH)
    print(SOURCE_PATH)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
